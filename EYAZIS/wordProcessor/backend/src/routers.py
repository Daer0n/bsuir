import io
import json
from docx import Document
from fastapi import APIRouter, UploadFile
from fastapi import HTTPException, status

from schemas import Word, Context
from functions import generate_word_form, get_lexemes, get_lexems_to_text, main_corpus, find_context

router = APIRouter(
    prefix="/word",
    tags=[""]
)

@router.post("/generate")
async def generate_word(word: Word) -> str:
    return generate_word_form(word)


@router.post('/parsing/sentence')
async def sentence_parsing(file: UploadFile):
    try:
        file_content = await file.read()
        file_object = io.BytesIO(file_content)
        doc = Document(file_object)
        sentence = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return get_lexems_to_text(get_lexemes(sentence))
    except FileNotFoundError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unprocessable entity")
    

@router.post('/save')
async def save(file: UploadFile):
    try:
        file_content = await file.read()
        file_object = io.BytesIO(file_content)
        doc = Document(file_object)
        sentence = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        lexemes = get_lexems_to_text(get_lexemes(sentence))
        
        with open('data/data.json', 'a') as json_file:
            for lexeme in lexemes:
                json_data = json.dumps(lexeme, ensure_ascii=False)
                json_file.write(json_data + '\n')
        
        return {'ok' : True}
    except FileNotFoundError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unprocessable entity")
    
@router.get('/base/search/{base}')
async def base_search(base: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if lexeme['Основа'] == base:
                data.append(lexeme)
    return data

@router.get('/part_of_speech/filter/{part_of_speech}')
async def part_of_speech_filter(part_of_speech: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if lexeme['ч.речи'] == part_of_speech:
                data.append(lexeme)
    return data

@router.get('/gender/filter/gender/{gender}')
async def gender_filter(gender: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if 'род' in lexeme and lexeme['род'] == gender:
                data.append(lexeme)
    return data

@router.get('/case/filter/case/{case}')
async def case_filter(case: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if 'падеж' in lexeme and lexeme['падеж'] == case:
                data.append(lexeme)
    return data

@router.get("/view/data")
async def view_data():
    with open('data/data.json', 'r') as json_file:
        file_content = json_file.read()
        return file_content
    
@router.get("/corpus/search/{sentence}")
async def search_corpus(sentence: str):
    return main_corpus(sentence)

@router.post("/search/context")
async def search_context(context: Context):
    return find_context(context.word, context.length, context.count)