import io
import json
from docx import Document
from fastapi import APIRouter, UploadFile
from fastapi import HTTPException, status

from schemas import Word
from functions import generate_word_form, get_lexemes, get_lexems_to_text

router = APIRouter(
    prefix="/word",
    tags=[""]
)

@router.post("/generate")
async def generate_word(word: Word) -> str:
    return generate_word_form(word)


@router.post('/parsing/sentence')
async def sentence_parsing(file: UploadFile):
    try:
        file_content = await file.read()
        file_object = io.BytesIO(file_content)
        doc = Document(file_object)
        sentence = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return get_lexems_to_text(get_lexemes(sentence))
    except FileNotFoundError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unprocessable entity")
    

@router.post('/save')
async def sentence_parsing(file: UploadFile):
    try:
        file_content = await file.read()
        file_object = io.BytesIO(file_content)
        doc = Document(file_object)
        sentence = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        lexemes = get_lexems_to_text(get_lexemes(sentence))
        
        with open('data/data.json', 'a') as json_file:
            for lexeme in lexemes:
                json_data = json.dumps(lexeme, ensure_ascii=False)
                json_file.write(json_data + '\n')
        
        return {'ok' : True}
    except FileNotFoundError:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unprocessable entity")
    
@router.post('/base/search')
async def base_search(base: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if lexeme['Основа'] == base:
                data.append(lexeme)
    return data

@router.post('/part/of/speech/search')
async def part_of_speech_search(base: str):
    with open('data/data.json') as json_file:
        data = []
        for line in json_file:
            lexeme = json.loads(line)
            if lexeme['ч.речи'] == base:
                data.append(lexeme)
    return data